"""
Word document processing utility for SnapQuote.

Converts Word document content to markdown format.
"""

from docx import Document
import io
import logging

logger = logging.getLogger(__name__)

def docx_to_markdown(docx_content: bytes) -> str:
    """
    Convert Word document content to markdown format.
    
    Args:
        docx_content (bytes): Word document file content as bytes
        
    Returns:
        str: Word document content converted to markdown format
    """
    try:
        # Create a BytesIO object from the docx content
        docx_file = io.BytesIO(docx_content)
        
        # Open the document
        doc = Document(docx_file)
        
        markdown_content = []
        markdown_content.append("# Word Document Content\n\n")
        
        # Process paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:  # Skip empty paragraphs
                continue
                
            # Check paragraph style to determine markdown formatting
            style_name = paragraph.style.name.lower() if paragraph.style else ""
            
            if 'heading 1' in style_name or 'title' in style_name:
                markdown_content.append(f"# {text}\n\n")
            elif 'heading 2' in style_name:
                markdown_content.append(f"## {text}\n\n")
            elif 'heading 3' in style_name:
                markdown_content.append(f"### {text}\n\n")
            elif 'heading 4' in style_name:
                markdown_content.append(f"#### {text}\n\n")
            elif 'heading 5' in style_name:
                markdown_content.append(f"##### {text}\n\n")
            elif 'heading 6' in style_name:
                markdown_content.append(f"###### {text}\n\n")
            else:
                # Regular paragraph
                markdown_content.append(f"{text}\n\n")
        
        # Process tables
        if doc.tables:
            markdown_content.append("## Tables\n\n")
            for i, table in enumerate(doc.tables):
                markdown_content.append(f"### Table {i + 1}\n\n")
                
                # Extract table data
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        # Clean cell text
                        cell_text = cell.text.strip().replace('\n', ' ').replace('\r', ' ')
                        row_data.append(cell_text)
                    table_data.append(row_data)
                
                if table_data:
                    # Create markdown table
                    if len(table_data) > 0:
                        # Header row
                        header = " | ".join(table_data[0])
                        separator = " | ".join(["---"] * len(table_data[0]))
                        markdown_content.append(f"| {header} |\n")
                        markdown_content.append(f"| {separator} |\n")
                        
                        # Data rows
                        for row in table_data[1:]:
                            row_text = " | ".join(row)
                            markdown_content.append(f"| {row_text} |\n")
                        
                        markdown_content.append("\n")
        
        if len(markdown_content) == 1:  # Only header, no content
            markdown_content.append("*[No readable text content found in this Word document]*\n")
        
        return "".join(markdown_content)
        
    except Exception as e:
        logger.error(f"Error processing Word document: {str(e)}")
        return f"# Word Document Processing Error\n\n*Error: {str(e)}*\n"